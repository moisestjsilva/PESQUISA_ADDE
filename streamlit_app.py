import streamlit as st
from docx import Document
from PIL import Image
import tempfile

def main():
    st.title('Imagens para Documento MNH')  # Alterado o título

    # Interface para upload de imagens
    uploaded_files = st.file_uploader("Selecione as imagens que deseja juntar", type=['jpg', 'png'], accept_multiple_files=True)

    # Interface para nomear o arquivo DOC
    doc_name = st.text_input("Nome do arquivo DOC", "meu_documento")

    # Botão para criar o documento
    if st.button("Criar Documento"):  # Alterado o nome do botão
        if uploaded_files:
            create_document(uploaded_files, doc_name)

    # Botão para limpar as seleções e o nome do documento
    if st.button("Limpar"):
        uploaded_files.clear()
        doc_name = "meu_documento"
        st.success("Seleção de imagens limpa.")

    # Mostra as imagens selecionadas
    if uploaded_files:
        st.write(f"{len(uploaded_files)} imagens selecionadas.")
        st.write("Lista de imagens:")
        for file in uploaded_files:
            st.write(file.name)

    # Mostra o nome do documento atual
    st.write(f"Nome do documento: {doc_name}")

def create_document(uploaded_files, doc_name):
    doc = Document()

    # Definir margens mínimas (em polegadas)
    margin_top = 0.5  # 0.5 polegadas
    margin_bottom = 0.5  # 0.5 polegadas
    margin_left = 0.5  # 0.5 polegadas
    margin_right = 0.5  # 0.5 polegadas

    # Converter polegadas para pontos (1 polegada = 72 pontos)
    margin_top_in_points = int(margin_top * 72)
    margin_bottom_in_points = int(margin_bottom * 72)
    margin_left_in_points = int(margin_left * 72)
    margin_right_in_points = int(margin_right * 72)

    # Configurar as margens do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = margin_top_in_points
        section.bottom_margin = margin_bottom_in_points
        section.left_margin = margin_left_in_points
        section.right_margin = margin_right_in_points

    # Calcular largura e altura máximas disponíveis no documento
    max_width = doc.sections[0].page_width - margin_left_in_points - margin_right_in_points
    max_height = doc.sections[0].page_height - margin_top_in_points - margin_bottom_in_points

    # Inicializa a largura e altura total das imagens na página
    total_width_on_page = 0
    total_height_on_page = 0

    # Lista para armazenar imagens na linha atual
    current_line_images = []

    for idx, file in enumerate(uploaded_files):
        # Abre a imagem e obtém suas dimensões
        img = Image.open(file)
        width, height = img.size

        # Calcula a proporção da imagem
        aspect_ratio = width / height

        # Calcula a nova largura e altura da imagem para ajustar ao documento
        new_width = max_width
        new_height = int(new_width / aspect_ratio)

        # Verifica se a imagem cabe na linha atual da página
        if total_height_on_page + new_height <= max_height:
            # Adiciona a imagem à linha atual
            current_line_images.append((file, new_width, new_height))
            total_width_on_page = max(total_width_on_page, new_width)
            total_height_on_page += new_height
        else:
            # Se não couber, adiciona uma nova linha de imagens
            add_images_to_document(doc, current_line_images)
            doc.add_page_break()

            # Reinicializa para uma nova linha de imagens
            current_line_images = [(file, new_width, new_height)]
            total_width_on_page = new_width
            total_height_on_page = new_height

        # Se for a última imagem, adiciona à linha atual
        if idx == len(uploaded_files) - 1:
            add_images_to_document(doc, current_line_images)

    # Salva o documento temporariamente
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc_path = tmp_file.name
        doc.save(doc_path)

    # Cria um botão de download para o arquivo DOC
    with open(doc_path, "rb") as file:
        st.download_button(
            label="Baixar Documento DOC",
            data=file,
            file_name=f"{doc_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    st.success(f"Documento {doc_name}.docx criado com sucesso!")

def add_images_to_document(doc, images):
    # Adiciona todas as imagens na lista à linha atual do documento
    for file, width, height in images:
        doc.add_picture(file, width=width, height=height)
        # Adiciona espaço entre imagens (opcional)
        doc.paragraphs[-1].runs[-1].add_break()

if __name__ == '__main__':
    main()
